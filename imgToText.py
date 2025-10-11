# -*- coding: utf-8 -*-
"""
Created on Thu Jan 30 15:55:17 2025
"""
#import standerd libraries
import math
import argparse
#import 3rd party libraries
from PIL import Image
import docx 
from docx.shared import RGBColor
#functions to load text/imgs
def load_photo(path):
    try:
        pillow_photo = Image.open(path)
        return [pillow_photo]
    except:
        return ['error']
def load_text(path):
    if path[0] == '"':
        new_path = path[1:-1]
    else:
        new_path = path
    try:
        file = open(new_path, 'r')
        text = file.read()
        file.close
        return text
    except:
        return 'error'
def split_string(string, x):
  """Splits a string into multiple strings of length x."""
  return [string[i:i + x] for i in range(0, len(string), x)]
def convert_img_to_text(text, image, truncate, repeat):
    photo = image
    text_string = text
    #get img size & calculate area
    width, height = photo.size
    photo_area = width * height
    #handle repeat flag
    if repeat == True:
        while len(text_string) < photo_area:
            text_string += text_string
    #handle truncate flag
    if truncate == True:
        print('Applying Truncation...')
        if len(text_string) > photo_area:
            difference = len(text_string) - photo_area
            text_string = text_string[:-difference]
    #get length of text
    text_length = len(text_string)
    #get ratio of the side lengths
    size_ratio = width / height
    #set up array for text
    text_array = []
    #find the side lengths of rectangle based on the side lengths ratio and area
    x = math.sqrt((text_length /  size_ratio))
    #change which will be bigger based on landscape or portrio
    text_width = round(size_ratio * x)
    text_height = round(x)
    #split the string into multiple strings the length of the text width
    text_array = split_string(text_string, text_width)
    #split each width into multiple lists with 1 charachter each
    for num in range(0, len(text_array)):
         split_text = split_string(text_array[num], 1)
         text_array[num] = [[x] for x in split_text]
    #find ratio of areas
    area_ratio = text_length / photo_area
    #find pixel values for each charachter
    print('Colouring Text...')
    if area_ratio > 1:
        #where there are more charachters then pixels
        #scale image
        scaled_image  = photo.resize((text_width, text_height))
        #move through each row of the text list
        for y in range(0, len(text_array)):
            #move through each charachter in the row
            for x in range(0, text_width):
                #get photo rgb for that corrdinent and append to the charchter list 
                try:
                    rgb_value = scaled_image.getpixel((x, y))
                    text_array[y][x].append(rgb_value)
                except:
                    break
    elif area_ratio == 1:
        #when there are the same number of charachters and pixels
        #move through each row of the text list
        for y in range(0, len(text_array)):
            #move through each charachter in the row
            for x in range(0, text_width):
                #get photo rgb for that corrdinent and append to the charchter list 
                rgb_value = photo.getpixel((x, y))
                text_array[y][x].append(rgb_value)
    else:
        #when there are less charachters than pixels
        #scale image
        scaled_image = photo.resize((math.ceil(text_width), math.ceil(text_height)))
        #move through each row of the text list
        for y in range(0, text_height):
            #move through each charachter in the row
            for x in range(0, text_width):
                #get photo rgb for that coordinent and append to the charchter list 
                rgb_value = scaled_image.getpixel((x, y))
                if rgb_value == (255, 255, 255, 255):
                    rgb_value = (250, 250, 250, 255)
                try:
                    text_array[y][x].append(rgb_value)
                except:
                    break
    return text_array
#function to create the word document
def create_document(text_array, word_doc_save_path):
    #create document
    print('Creating Document...')
    doc = docx.Document() 
    #add paragraph to document
    para = doc.add_paragraph()
    num_lines = len(text_array)
    for line, list in enumerate(text_array):
        #move to a new line to keep img intact
        para.add_run('\n')
        num_chars = len(list)
        for char_num, charachter in enumerate(list):
            print(f'char {char_num}/{num_chars}, line {line}/{num_lines}')
            #iterate through charachters in a line
            if len(charachter) == 2:
                #select colour for chrahcter
                if charachter[1] == (255, 255, 255, 255):
                    colour = RGBColor(254, 254, 254)
                else:
                    colour = RGBColor(charachter[1][0], charachter[1][1], charachter[1][2])
            else:
                colour = RGBColor(0, 0, 0)
            #add charachter to paragraph
            run = para.add_run(charachter[0])
            #change chrachter colour to be the loaded colour
            run.font.color.rgb = colour
    #save docx file
    print('Saving File...')
    doc.save(word_doc_save_path)

def run(photo_path, text_path, word_doc_save_path, truncate, repeat):
    #load photo
    print('Loading Image...')
    photos = load_photo(photo_path)
    #get text/load txt file
    print('Loading Text...')
    text = load_text(text_path)
    #create array of lines, charachters, and colours
    text_array = convert_img_to_text(text, photos[0], truncate, repeat)
    #create & save word document using array created above
    create_document(text_array, word_doc_save_path)

#set up argparse parameters
parser = argparse.ArgumentParser(prog='imgToText', description='Convert images to word document')
parser.add_argument('imagePath', help='path to image file') 
parser.add_argument('textPath', help='path to text document') 
parser.add_argument('outputPath', help='path to output word document to')
parser.add_argument('-t', '--truncate', action='store_true', help='truncate text if the image is larger')
parser.add_argument('-r', '--repeat', action='store_true', help='repeat text until it is the size of the image')
#get arguments
args = parser.parse_args()
run(args.imagePath, args.textPath, args.outputPath, args.truncate, args.repeat)
print('Done!')